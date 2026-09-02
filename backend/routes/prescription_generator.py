import os

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from docx import Document


def generate_prescription_files(prescription):

    base_dir = os.path.dirname(
        os.path.dirname(
            os.path.abspath(__file__)
        )
    )

    output_dir = os.path.join(
        base_dir,
        "prescriptions"
    )

    os.makedirs(
        output_dir,
        exist_ok=True
    )

    prescription_id = prescription["id"]

    pdf_filename = f"{prescription_id}.pdf"
    docx_filename = f"{prescription_id}.docx"

    pdf_path = os.path.join(
        output_dir,
        pdf_filename
    )

    docx_path = os.path.join(
        output_dir,
        docx_filename
    )

    generate_pdf(
        prescription,
        pdf_path
    )

    generate_docx(
        prescription,
        docx_path
    )

    return {
        "pdf": f"prescriptions/{pdf_filename}",
        "docx": f"prescriptions/{docx_filename}"
    }


# =====================================================
# PDF
# =====================================================

def generate_pdf(prescription, output_path):

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "PrescriptionTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=20,
        spaceAfter=5
    )

    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=10,
        textColor=colors.grey
    )

    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=6
    )

    normal_style = styles["Normal"]

    document = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    story = []

    # -------------------------------------------------
    # HEADER
    # -------------------------------------------------

    story.append(
        Paragraph(
            "MEDIBRIDGE",
            title_style
        )
    )

    story.append(
        Paragraph(
            "MEDICAL PRESCRIPTION",
            subtitle_style
        )
    )

    story.append(
        Spacer(1, 20)
    )

    # -------------------------------------------------
    # DOCTOR / PATIENT INFORMATION
    # -------------------------------------------------

    doctor_info = [
        [
            Paragraph("<b>Doctor</b>", normal_style),
            prescription.get("doctor_name") or ""
        ],
        [
            Paragraph("<b>Specialization</b>", normal_style),
            prescription.get("specialization") or ""
        ],
        [
            Paragraph("<b>Patient</b>", normal_style),
            prescription.get("patient_name") or ""
        ],
        [
            Paragraph("<b>Date</b>", normal_style),
            prescription.get("date") or ""
        ]
    ]

    info_table = Table(
        doctor_info,
        colWidths=[120, 350]
    )

    info_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 7)
        ])
    )

    story.append(info_table)

    # -------------------------------------------------
    # DIAGNOSIS
    # -------------------------------------------------

    story.append(
        Paragraph(
            "DIAGNOSIS",
            heading_style
        )
    )

    story.append(
        Paragraph(
            str(prescription.get("diagnosis") or ""),
            normal_style
        )
    )

    # -------------------------------------------------
    # MEDICATION
    # -------------------------------------------------

    story.append(
        Paragraph(
            "MEDICATION",
            heading_style
        )
    )

    medicine_rows = [
        [
            "Medicine",
            "Dosage",
            "Frequency",
            "Duration",
            "Instructions"
        ]
    ]

    medicines = prescription.get("medicines") or []

    for medicine in medicines:

        medicine_rows.append([
            str(medicine.get("name") or ""),
            str(medicine.get("dosage") or ""),
            str(medicine.get("frequency") or ""),
            str(medicine.get("duration") or ""),
            str(medicine.get("instructions") or "")
        ])

    medicine_table = Table(
        medicine_rows,
        colWidths=[100, 75, 80, 65, 100],
        repeatRows=1
    )

    medicine_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 5)
        ])
    )

    story.append(medicine_table)

    # -------------------------------------------------
    # ADDITIONAL ADVICE
    # -------------------------------------------------

    story.append(
        Paragraph(
            "ADDITIONAL INSTRUCTIONS",
            heading_style
        )
    )

    story.append(
        Paragraph(
            str(prescription.get("advice") or ""),
            normal_style
        )
    )

    # -------------------------------------------------
    # FOLLOW UP
    # -------------------------------------------------

    story.append(
        Paragraph(
            "FOLLOW-UP",
            heading_style
        )
    )

    story.append(
        Paragraph(
            str(
                prescription.get(
                    "follow_up_date"
                ) or "Not specified"
            ),
            normal_style
        )
    )

    story.append(
        Spacer(1, 30)
    )

    story.append(
        Paragraph(
            "Doctor Signature",
            normal_style
        )
    )

    story.append(
        Spacer(1, 25)
    )

    story.append(
        Paragraph(
            str(prescription.get("doctor_name") or ""),
            normal_style
        )
    )

    document.build(story)


# =====================================================
# DOCX
# =====================================================

def generate_docx(prescription, output_path):

    document = Document()

    # -------------------------------------------------
    # TITLE
    # -------------------------------------------------

    title = document.add_heading(
        "MEDIBRIDGE",
        level=0
    )

    title.alignment = 1

    subtitle = document.add_paragraph(
        "MEDICAL PRESCRIPTION"
    )

    subtitle.alignment = 1

    document.add_paragraph("")

    # -------------------------------------------------
    # DOCTOR / PATIENT INFORMATION
    # -------------------------------------------------

    table = document.add_table(
        rows=4,
        cols=2
    )

    table.style = "Table Grid"

    information = [
        (
            "Doctor",
            str(prescription.get("doctor_name") or "")
        ),
        (
            "Specialization",
            str(prescription.get("specialization") or "")
        ),
        (
            "Patient",
            str(prescription.get("patient_name") or "")
        ),
        (
            "Date",
            str(prescription.get("date") or "")
        )
    ]

    for row, (label, value) in zip(
        table.rows,
        information
    ):

        row.cells[0].text = label
        row.cells[1].text = value

    # -------------------------------------------------
    # DIAGNOSIS
    # -------------------------------------------------

    document.add_heading(
        "Diagnosis",
        level=2
    )

    document.add_paragraph(
        str(prescription.get("diagnosis") or "")
    )

    # -------------------------------------------------
    # MEDICATION
    # -------------------------------------------------

    document.add_heading(
        "Medication",
        level=2
    )

    medicines = prescription.get("medicines") or []

    medicine_table = document.add_table(
        rows=1,
        cols=5
    )

    medicine_table.style = "Table Grid"

    headers = [
        "Medicine",
        "Dosage",
        "Frequency",
        "Duration",
        "Instructions"
    ]

    for i, header in enumerate(headers):

        medicine_table.rows[0].cells[i].text = header

    for medicine in medicines:

        row = medicine_table.add_row()

        row.cells[0].text = str(
            medicine.get("name") or ""
        )

        row.cells[1].text = str(
            medicine.get("dosage") or ""
        )

        row.cells[2].text = str(
            medicine.get("frequency") or ""
        )

        row.cells[3].text = str(
            medicine.get("duration") or ""
        )

        row.cells[4].text = str(
            medicine.get("instructions") or ""
        )

    # -------------------------------------------------
    # ADDITIONAL INSTRUCTIONS
    # -------------------------------------------------

    document.add_heading(
        "Additional Instructions",
        level=2
    )

    document.add_paragraph(
        str(prescription.get("advice") or "")
    )

    # -------------------------------------------------
    # FOLLOW UP
    # -------------------------------------------------

    document.add_heading(
        "Follow-up",
        level=2
    )

    document.add_paragraph(
        str(
            prescription.get(
                "follow_up_date"
            ) or "Not specified"
        )
    )

    document.add_paragraph("")

    document.add_paragraph(
        "Doctor Signature"
    )

    document.add_paragraph(
        str(
            prescription.get(
                "doctor_name"
            ) or ""
        )
    )

    document.save(output_path)